from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
import os
import win32com.client as win32
import pythoncom
from docx.shared import Pt

app = Flask(__name__)

# 엑셀 파일 로드
df = pd.read_excel('data.xlsx')

pythoncom.CoInitialize()
# 수료증 템플릿 로드
template_docx_path = 'certificate.docx'
doc = Document(template_docx_path)

def convert_docx_to_pdf(input_path, output_path):
    
    # Word 객체 생성
    pythoncom.CoInitialize()
    base_path = "C:\\Users\\Administrator\\Desktop\\auto_certificate\\"
    word = win32.gencache.EnsureDispatch("Word.Application")
    doc = word.Documents.Open(base_path + input_path)
    
    # PDF로 저장
    doc.SaveAs(base_path + output_path, FileFormat=17)  # 17은 PDF 형식
    
    # 닫기
    doc.Close()
    word.Quit()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_certificate', methods=['POST'])
def generate_certificate():
    name = request.form['name']
    email = request.form['email']

    # 이름과 이메일 확인
    match = df[(df['이름'] == name) & (df['이메일'] == email)]

    if not match.empty:
        new_doc = Document(template_docx_path)
        for p in new_doc.paragraphs:
            if '{{name}}' in p.text:
                for run in p.runs:
                    run.text = run.text.replace('{{name}}', name)
                    run.font.size = Pt(17)
                    run.font.bold = True
        
        # .docx로 저장
        temp_path = f"temp_{name}.docx"
        new_doc.save(temp_path)

        pdf_path = f"{name}_knockOn수료증.pdf"
        convert_docx_to_pdf(temp_path, pdf_path)

        # 임시 파일 삭제
        os.remove(temp_path)

        return send_file(pdf_path, as_attachment=True)
    else:
        return "이름 또는 이메일이 잘못되었습니다."

if __name__ == '__main__':
    app.run(debug=False, host="0.0.0.0")
